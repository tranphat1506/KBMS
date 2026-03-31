import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import markdown
from bs4 import BeautifulSoup

# --- Configuration ---
DOCS_DIR = 'docs'
OUTPUT_FILE = 'DATN_LeChauTranPhat_KBMS.docx'
FONT_NAME = 'Times New Roman'
FONT_SIZE = Pt(13)

def set_font(run):
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)

def add_heading(doc, text, level):
    h = doc.add_heading(text, level)
    for run in h.runs:
        set_font(run)
    return h

def add_paragraph(doc, text="", style='Normal', is_latex=False):
    p = doc.add_paragraph(style=style)
    # Basic LaTeX handling: Remove $$ and bold the content
    if '$$' in text:
        parts = re.split(r'(\$\$.*?\$\$)', text)
        for part in parts:
            if part.startswith('$$') and part.endswith('$$'):
                content = part[2:-2].strip()
                run = p.add_run(content)
                run.bold = True
                run.italic = True
            else:
                run = p.add_run(part)
            set_font(run)
    else:
        run = p.add_run(text)
        set_font(run)
    return p

def create_cover_page(doc):
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p1.add_run("TRƯỜNG ĐẠI HỌC QUỐC TẾ HỒNG BÀNG\nBỘ MÔN CÔNG NGHỆ THÔNG TIN")
    run.bold = True
    run.font.size = Pt(14)
    set_font(run)
    for _ in range(5): doc.add_paragraph()
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("ĐỒ ÁN TỐT NGHIỆP")
    run.bold = True
    run.font.size = Pt(28)
    set_font(run)
    for _ in range(2): doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p3.add_run("ĐỀ TÀI: “THIẾT KẾ HỆ HỖ TRỢ QUẢN TRỊ \nCƠ SỞ TRI THỨC DẠNG COKB”")
    run.bold = True
    run.font.size = Pt(20)
    set_font(run)
    for _ in range(6): doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p4.paragraph_format.left_indent = Inches(1.5)
    run = p4.add_run("GVHD: Gs. ĐỖ VĂN NHƠN & Ths. MAI TRUNG THÀNH\nSVTH: LÊ CHÂU TRẦN PHÁT - 2211110068 - TH22DH-CN2")
    run.bold = True
    run.font.size = Pt(13)
    set_font(run)
    for _ in range(8): doc.add_paragraph()
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p5.add_run("TP. Hồ Chí Minh, 2026")
    run.bold = True
    run.font.size = Pt(13)
    set_font(run)
    doc.add_page_break()

def add_toc_placeholder(doc):
    add_heading(doc, "MỤC LỤC", level=0)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._element.append(fldChar)
    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._element.append(instrText)
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'separate')
    run._element.append(fldChar)
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar)
    doc.add_page_break()

def strip_leading_numbers(text):
    # Removes patterns like "01.1. ", "1.2. ", etc.
    return re.sub(r'^[\d\s\.]+', '', text).strip()

def clean_label_numbering(text):
    # Transforms "Hình 04.1:" into "Hình 4.1:"
    def replacer(match):
        label = match.group(1)
        number = match.group(2).lstrip('0')
        if not number or number.startswith('.'): number = '0' + number
        number = re.sub(r'\.0(\d)', r'.\1', number) # 4.01 -> 4.1
        rest = match.group(3)
        return f"{label} {number}{rest}"
    
    return re.sub(r'(Hình|Bảng)\s+([\d\.]+)(.*)', replacer, text)

def process_markdown(doc, md_content, chapter_idx, file_idx, folder_name):
    # Pre-clean LaTeX symbols for paragraphs
    md_content = md_content.replace('\\$', '$')

    html = markdown.markdown(md_content, extensions=['tables', 'fenced_code'])
    soup = BeautifulSoup(html, 'html.parser')
    
    h2_counter = 0
    h3_counter = 0

    for tag in soup.find_all(recursive=False):
        if tag.name == 'h1':
            title = strip_leading_numbers(tag.get_text())
            add_heading(doc, f"{chapter_idx}.{file_idx}. {title}", level=1)
            h2_counter = 0
        elif tag.name == 'h2':
            h2_counter += 1
            h3_counter = 0
            title = strip_leading_numbers(tag.get_text())
            add_heading(doc, f"{chapter_idx}.{file_idx}.{h2_counter}. {title}", level=2)
        elif tag.name == 'h3':
            h3_counter += 1
            title = strip_leading_numbers(tag.get_text())
            add_heading(doc, f"{chapter_idx}.{file_idx}.{h2_counter}.{h3_counter}. {title}", level=3)
        elif tag.name == 'p':
            text = tag.get_text().strip()
            if text.startswith('*Hình') or text.startswith('*Bảng'):
                # Clean leading zero in label
                clean_text = clean_label_numbering(text.replace('*', ''))
                p = add_paragraph(doc, clean_text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif tag.find('img'):
                img_tag = tag.find('img')
                img_src = img_tag.get('src')
                if img_src.startswith('../'):
                    img_path = os.path.join(DOCS_DIR, img_src[3:])
                else:
                    img_path = os.path.join(DOCS_DIR, folder_name, img_src)
                
                if os.path.exists(img_path):
                    doc.add_picture(img_path, width=Inches(5.5))
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                add_paragraph(doc, text)
        elif tag.name in ['ul', 'ol']:
            for li in tag.find_all('li'):
                p = add_paragraph(doc, f"• {li.get_text().strip()}", style='List Bullet')
        elif tag.name == 'table':
            rows = tag.find_all('tr')
            if not rows: continue
            cols_count = len(rows[0].find_all(['td', 'th']))
            table = doc.add_table(rows=len(rows), cols=cols_count)
            table.style = 'Table Grid'
            for i, row in enumerate(rows):
                cells = row.find_all(['td', 'th'])
                for j, cell in enumerate(cells):
                    table.cell(i, j).text = cell.get_text().strip()
            doc.add_paragraph()
        elif tag.name == 'pre':
            text = tag.get_text().strip()
            p = add_paragraph(doc, text)
            p.style = doc.styles['No Spacing']
            for run in p.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)

def main():
    print("Regenerating: Final Formatting Fixes...")
    
    lof_entries = []
    lot_entries = []
    folders = sorted([f for f in os.listdir(DOCS_DIR) if re.match(r'^\d{2}-', f)])
    
    # 1. Collect LOF/LOT
    for folder in folders:
        folder_path = os.path.join(DOCS_DIR, folder)
        files = sorted([f for f in os.listdir(folder_path) if f.endswith('.md')])
        for file in files:
            with open(os.path.join(folder_path, file), 'r', encoding='utf-8') as f:
                content = f.read()
                # Use more robust regex
                figures = re.findall(r'\*Hình\s+[\d\.]+:.*?\*', content)
                for fig in figures:
                    lof_entries.append(clean_label_numbering(fig.replace('*', '')))
                tables = re.findall(r'\*Bảng\s+[\d\.]+:.*?\*', content)
                for tab in tables:
                    lot_entries.append(clean_label_numbering(tab.replace('*', '')))

    # 2. Build Document
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = FONT_SIZE
    
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(0.8)

    create_cover_page(doc)
    add_toc_placeholder(doc)
    
    add_heading(doc, "DANH MỤC HÌNH ẢNH", level=1)
    for entry in lof_entries: add_paragraph(doc, entry)
    doc.add_page_break()
    
    add_heading(doc, "DANH MỤC BẢNG", level=1)
    for entry in lot_entries: add_paragraph(doc, entry)
    doc.add_page_break()

    # 3. Add Content
    for folder in folders:
        chapter_idx = int(folder.split('-')[0])
        folder_path = os.path.join(DOCS_DIR, folder)
        files = sorted([f for f in os.listdir(folder_path) if f.endswith('.md')])
        for f_idx, file in enumerate(files, 1):
            with open(os.path.join(folder_path, file), 'r', encoding='utf-8') as f:
                process_markdown(doc, f.read(), chapter_idx, f_idx, folder)
        doc.add_page_break()

    doc.save(OUTPUT_FILE)
    print(f"Success! Report regenerated: {OUTPUT_FILE}")
    print(f"Stats: {len(lof_entries)} Fig, {len(lot_entries)} Tab.")

if __name__ == '__main__':
    main()

if __name__ == '__main__':
    main()
